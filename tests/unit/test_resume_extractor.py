from io import BytesIO

import pytest
from docx import Document

from app.tools.resume_extractor import (
    ResumeExtractionError,
    extract_resume_text,
)


def test_extract_txt_resume_success() -> None:
    file_bytes = b"John Doe\nPython Developer\nFastAPI, PostgreSQL"
    text = extract_resume_text("resume.txt", file_bytes)

    assert "John Doe" in text
    assert "Python Developer" in text


def test_extract_txt_resume_invalid_utf8() -> None:
    file_bytes = b"\xff\xfe\x00\x00"

    with pytest.raises(ResumeExtractionError, match="Failed to decode TXT file"):
        extract_resume_text("resume.txt", file_bytes)


def test_extract_docx_resume_success() -> None:
    doc = Document()
    doc.add_paragraph("Jane Doe")
    doc.add_paragraph("Backend Engineer")
    doc.add_paragraph("Skills: Python, SQLAlchemy")

    buffer = BytesIO()
    doc.save(buffer)

    text = extract_resume_text("resume.docx", buffer.getvalue())

    assert "Jane Doe" in text
    assert "Backend Engineer" in text
    assert "Python" in text


def test_reject_unsupported_file_type() -> None:
    file_bytes = b"some content"

    with pytest.raises(ResumeExtractionError, match="Unsupported file type"):
        extract_resume_text("resume.exe", file_bytes)


def test_reject_empty_extracted_text() -> None:
    file_bytes = b"   \n\t   "

    with pytest.raises(ResumeExtractionError, match="No readable text could be extracted"):
        extract_resume_text("resume.txt", file_bytes)